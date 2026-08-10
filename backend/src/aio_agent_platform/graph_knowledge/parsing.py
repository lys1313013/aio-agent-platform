"""Document text extraction for graph knowledge base uploads (文档解析).

Extracts plain text / Markdown from uploaded files so they can be chunked
and fed to the entity/relation extraction pipeline. Supported formats:
.md/.markdown/.txt, .html/.htm, .pdf, .docx.
"""

from __future__ import annotations

import io

SUPPORTED_EXTENSIONS = (".md", ".markdown", ".txt", ".html", ".htm", ".pdf", ".docx")


class DocumentParseError(ValueError):
    """Raised when an uploaded document cannot be parsed (用户可读的错误信息)."""


class ScannedPDFError(DocumentParseError):
    """Raised when a PDF has no extractable text layer — needs OCR (MinerU)."""


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_html(data: bytes) -> str:
    from markdownify import markdownify

    return markdownify(_decode_text(data))


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:
        raise DocumentParseError(f"PDF 解析失败:{e}") from e
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(p for p in pages if p)
    if not text:
        raise ScannedPDFError("PDF 无文本层(扫描件),需要通过 MinerU OCR 解析")
    return text


def _extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise DocumentParseError(f"Word 文档解析失败:{e}") from e
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts)
    if not text:
        raise DocumentParseError("Word 文档未提取到文本内容")
    return text


def extract_text(filename: str, data: bytes) -> str:
    """Extract text from uploaded file bytes, dispatching by extension.

    Raises DocumentParseError for unsupported formats or parse failures.
    """
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in (".md", ".markdown", ".txt"):
        return _decode_text(data)
    if ext in (".html", ".htm"):
        return _extract_html(data)
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)
    supported = "、".join(SUPPORTED_EXTENSIONS)
    raise DocumentParseError(f"不支持的文件格式 '{ext or '(无扩展名)'}',支持:{supported}")
